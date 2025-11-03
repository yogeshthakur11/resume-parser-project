"""
Resume Parser API using FastAPI and Groq (Llama 3.2)
Technical Assessment Solution with Chain-of-Thought Prompting
With Proper HTTP Status Codes
"""

from fastapi import FastAPI, File, UploadFile, HTTPException, status
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import Optional, List
import os
import json
from pathlib import Path
import tempfile
from dotenv import load_dotenv

# Document processing libraries
import PyPDF2
import docx
from groq import Groq

# Initialize FastAPI app
app = FastAPI(
    title="AI Resume Parser",
    description="Extract structured information from resumes using Transformer models",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Groq client
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

# ============================================
# Pydantic Models for Response Structure
# ============================================

class ContactInfo(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    location: Optional[str] = None

class Education(BaseModel):
    institution: Optional[str] = None
    degree: Optional[str] = None
    field_of_study: Optional[str] = None
    graduation_year: Optional[str] = None
    gpa: Optional[str] = None

class WorkExperience(BaseModel):
    company: Optional[str] = None
    position: Optional[str] = None
    duration: Optional[str] = None
    description: Optional[str] = None
    location: Optional[str] = None

class Certification(BaseModel):
    name: Optional[str] = None
    issuer: Optional[str] = None
    date: Optional[str] = None

class Project(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    technologies: Optional[List[str]] = None
    link: Optional[str] = None

class ResumeData(BaseModel):
    is_resume: bool = True
    not_resume_reason: Optional[str] = None
    filename: Optional[str] = None
    contact_info: ContactInfo
    education: List[Education] = []
    work_experience: List[WorkExperience] = []
    skills: List[str] = []
    certifications: List[Certification] = []
    projects: List[Project] = []
    summary: Optional[str] = None

# ============================================
# File Processing Functions
# ============================================

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, 
            detail=f"Error reading PDF: {str(e)}"
        )

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file with enhanced extraction for headers, footers, and text boxes"""
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        # 1. Extract from HEADERS (where names often appear)
        for section in doc.sections:
            # Header extraction
            header = section.header
            for paragraph in header.paragraphs:
                para_text = paragraph.text.strip()
                if para_text and not para_text.startswith('![](media/'):
                    text_parts.append(para_text)
            
            # Footer extraction (sometimes contact info is here)
            footer = section.footer
            for paragraph in footer.paragraphs:
                para_text = paragraph.text.strip()
                if para_text and not para_text.startswith('![](media/'):
                    text_parts.append(para_text)
        
        # 2. Extract from main document paragraphs
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text and not para_text.startswith('![](media/'):
                text_parts.append(para_text)
        
        # 3. Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and not cell_text.startswith('![](media/'):
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        # 4. Extract from text boxes and shapes (requires python-docx with XML access)
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import qn
            
            # Access the underlying XML to find text boxes
            for element in doc.element.body.iter():
                # Look for text box content
                if element.tag.endswith('txbxContent'):
                    for paragraph in element.iter():
                        if paragraph.tag.endswith('t'):  # Text element
                            text_content = paragraph.text
                            if text_content and text_content.strip():
                                text_parts.append(text_content.strip())
        except Exception as xml_error:
            # If XML extraction fails, continue without it
            print(f"XML extraction skipped: {xml_error}")
        
        # Join all parts with newlines
        text = "\n".join(text_parts)
        
        # Clean up the text
        text = clean_extracted_text(text)
        
        return text.strip()
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, 
            detail=f"Error reading DOCX: {str(e)}"
        )


def clean_extracted_text(text: str) -> str:
    """Remove noise from extracted text"""
    import re
    
    # Remove image references
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'\{width=.*?\}', '', text)
    text = re.sub(r'\{height=.*?\}', '', text)
    
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    
    # Remove markdown artifacts
    text = re.sub(r'#+\s*', '', text)
    
    return text

def validate_file_format(filename: str) -> bool:
    """Validate if file format is supported"""
    if not filename:
        return False
    
    allowed_extensions = ['.pdf', '.docx', '.doc']
    file_ext = Path(filename).suffix.lower()
    
    return file_ext in allowed_extensions

def extract_text_from_file(file_path: str, filename: str) -> str:
    """Route to appropriate text extraction based on file type"""
    file_ext = Path(filename).suffix.lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, 
            detail=f"Unsupported file format: {file_ext}. Allowed formats: PDF, DOCX, DOC"
        )

# ============================================
# Chain-of-Thought Prompt Engineering
# ============================================

def create_cot_prompt(resume_text: str) -> str:
    """
    Create a Chain-of-Thought prompt for resume parsing with format detection
    """
    
    prompt = f"""You are an expert resume parser AI. Your task is to:
1. First determine if this is actually a resume/CV
2. If yes, extract structured information systematically

**Chain-of-Thought Process:**

Step 0: VALIDATE IF THIS IS A RESUME
- Check for resume indicators: contact info, education, work experience, skills
- If missing ALL key sections → NOT a resume
- Common non-resume documents: invoices, reports, letters, articles, books

Step 1: IDENTIFY CONTACT INFORMATION
- Scan for name (usually at the top)
- Look for email addresses (contains @ symbol)
- Find phone numbers (various formats: +91, (xxx) xxx-xxxx, etc.)
- Locate LinkedIn profile URLs
- Identify location/address if present

Step 2: EXTRACT EDUCATION HISTORY
- Find education section (keywords: Education, Academic, Qualification)
- For each entry, extract:
  * Institution/University name
  * Degree type (Bachelor's, Master's, PhD, etc.)
  * Field of study/Major
  * Graduation year or expected graduation
  * GPA/Percentage if mentioned

Step 3: ANALYZE WORK EXPERIENCE
- Locate work experience section (keywords: Experience, Employment, Work History)
- For each role, extract:
  * Company name
  * Job title/Position
  * Duration (start date - end date or "Present")
  * Key responsibilities and achievements
  * Location if mentioned

Step 4: IDENTIFY SKILLS
- Find skills section (keywords: Skills, Technical Skills, Core Competencies)
- Extract both technical and soft skills
- Include programming languages, tools, frameworks, methodologies

Step 5: EXTRACT CERTIFICATIONS
- Look for certifications section
- Extract certificate name, issuing organization, and date

Step 6: FIND PROJECTS
- Identify projects section
- For each project: name, description, technologies used, links

Step 7: CAPTURE PROFESSIONAL SUMMARY
- Look for summary/objective section at the beginning
- Extract the professional summary or career objective

**DOCUMENT TEXT:**
{resume_text}

**OUTPUT INSTRUCTIONS:**
Now, following the above steps, provide ONLY a valid JSON response in this exact format (no additional text):

{{
  "is_resume": true or false,
  "not_resume_reason": "Brief explanation if not a resume, otherwise null",
  "contact_info": {{
    "name": "Full Name or null",
    "email": "email@example.com or null",
    "phone": "+1234567890 or null",
    "linkedin": "linkedin.com/in/profile or null",
    "location": "City, Country or null"
  }},
  "education": [
    {{
      "institution": "University Name",
      "degree": "Bachelor's/Master's",
      "field_of_study": "Computer Science",
      "graduation_year": "2023",
      "gpa": "3.8/4.0"
    }}
  ],
  "work_experience": [
    {{
      "company": "Company Name",
      "position": "Job Title",
      "duration": "Jan 2020 - Dec 2022",
      "description": "Key responsibilities and achievements",
      "location": "City, Country"
    }}
  ],
  "skills": ["Python", "Machine Learning", "FastAPI"],
  "certifications": [
    {{
      "name": "Certification Name",
      "issuer": "Issuing Organization",
      "date": "2023"
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "description": "Brief description",
      "technologies": ["Tech1", "Tech2"],
      "link": "github.com/project"
    }}
  ],
  "summary": "Professional summary or objective"
}}

CRITICAL: If document is NOT a resume, set "is_resume": false and provide reason. Still return the JSON structure with empty/null values for other fields.

Return ONLY the JSON object, no other text."""

    return prompt

# ============================================
# LLM Processing Function
# ============================================

def parse_resume_with_llm(resume_text: str) -> dict:
    """
    Use Groq API (Llama 3.3) to parse resume with chain-of-thought prompting
    """
    try:
        # Create chain-of-thought prompt
        prompt = create_cot_prompt(resume_text)
        
        # Call Groq API
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert resume parser that extracts structured information from resumes. Always respond with valid JSON only, no additional text."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=2000,
            top_p=0.9
        )
        
        # Extract response
        response_text = chat_completion.choices[0].message.content.strip()
        
        # Clean response (remove markdown code blocks if present)
        if response_text.startswith("```json"):
            response_text = response_text.replace("```json", "").replace("```", "").strip()
        elif response_text.startswith("```"):
            response_text = response_text.replace("```", "").strip()
        
        # Parse JSON
        parsed_data = json.loads(response_text)
        
        return parsed_data
        
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, 
            detail=f"Failed to parse LLM response as JSON: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, 
            detail=f"Error in LLM processing: {str(e)}"
        )

# ============================================
# API Endpoints
# ============================================

@app.get("/", status_code=status.HTTP_200_OK)
async def root():
    """Root endpoint with API information"""
    return {
        "message": "AI Resume Parser API",
        "version": "1.0.0",
        "endpoints": {
            "/parse-resume": "POST - Upload and parse resume file",
            "/health": "GET - Health check"
        },
        "supported_formats": ["PDF", "DOCX", "DOC"],
        "model": "llama-3.3-70b-versatile",
        "status_codes": {
            "200": "Success - All files parsed successfully",
            "207": "Multi-Status - Some files succeeded, some failed",
            "400": "Bad Request - All files failed or invalid request",
            "415": "Unsupported Media Type - Invalid file format",
            "422": "Unprocessable Entity - File content extraction error",
            "500": "Internal Server Error - Server processing error"
        }
    }

@app.get("/health", status_code=status.HTTP_200_OK)
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy", 
        "service": "resume-parser",
        "version": "1.0.0"
    }

@app.post("/parse-resume")
async def parse_resume(files: List[UploadFile] = File(...)):
    """
    Main endpoint to parse single or multiple resume files
    
    Args:
        files: One or more resume files (PDF, DOCX, or DOC format)
    
    Returns:
        JSON with extracted resume information
        
    Status Codes:
        - 200: All files successfully parsed
        - 207: Partial success (some succeeded, some failed)
        - 400: All files failed or no files provided
        - 415: Unsupported file format
        - 422: File content extraction error
        - 500: Internal server error
    """
    
    # Validate that files were provided
    if not files or len(files) == 0:
        return JSONResponse(
            content={
                "status": "error",
                "message": "No files provided",
                "total_files": 0,
                "successful": 0,
                "failed": 0,
                "results": None,
                "errors": None
            },
            status_code=status.HTTP_400_BAD_REQUEST
        )
    
    # Handle single file vs multiple files
    results = []
    errors = []
    
    for idx, file in enumerate(files):
        temp_file_path = None
        try:
            # Validate file format
            if not validate_file_format(file.filename):
                errors.append({
                    "file": file.filename,
                    "error": "Unsupported file format. Allowed: PDF, DOCX, DOC",
                    "status_code": status.HTTP_415_UNSUPPORTED_MEDIA_TYPE
                })
                continue
            
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            # Extract text from file
            resume_text = extract_text_from_file(temp_file_path, file.filename)
            
            if not resume_text or len(resume_text.strip()) < 50:
                errors.append({
                    "file": file.filename,
                    "error": "Could not extract sufficient text from the file",
                    "status_code": status.HTTP_422_UNPROCESSABLE_ENTITY
                })
                continue
            
            # Parse resume using LLM with Chain-of-Thought
            parsed_data = parse_resume_with_llm(resume_text)
            
            # Check if it's actually a resume
            if not parsed_data.get("is_resume", True):
                errors.append({
                    "file": file.filename,
                    "error": "Uploaded file does not appear to be a resume",
                    "reason": parsed_data.get("not_resume_reason", "Unknown"),
                    "status_code": status.HTTP_422_UNPROCESSABLE_ENTITY
                })
                continue
            
            # Add filename to result
            parsed_data["filename"] = file.filename
            results.append(parsed_data)
            
        except HTTPException as http_exc:
            errors.append({
                "file": file.filename,
                "error": http_exc.detail,
                "status_code": http_exc.status_code
            })
        except Exception as e:
            errors.append({
                "file": file.filename,
                "error": f"Unexpected error: {str(e)}",
                "status_code": status.HTTP_500_INTERNAL_SERVER_ERROR
            })
        finally:
            # Cleanup temporary file
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception as cleanup_error:
                    print(f"Warning: Could not delete temp file: {cleanup_error}")
    
    # Determine status code and message based on results
    if len(results) == 0 and len(errors) > 0:
        # All files failed
        response_status_code = status.HTTP_400_BAD_REQUEST
        response_message = "All files failed to process"
    elif len(errors) > 0 and len(results) > 0:
        # Partial success
        response_status_code = status.HTTP_207_MULTI_STATUS
        response_message = "Partial success - some files processed successfully"
    else:
        # All files succeeded
        response_status_code = status.HTTP_200_OK
        response_message = "All files processed successfully"
    
    # Build response
    response = {
        "status": "success" if len(results) > 0 else "error",
        "message": response_message,
        "total_files": len(files),
        "successful": len(results),
        "failed": len(errors),
        "results": results if len(files) > 1 else (results[0] if results else None),
        "errors": errors if errors else None
    }
    
    return JSONResponse(content=response, status_code=response_status_code)

# ============================================
# Run Application
# ============================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=8000,
        log_level="info"
    )